"""将 项目实现详解.md 转换为可打印的 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading_elm.append(shading)

def add_code_block(doc, code_text):
    """添加代码块（灰底等宽字体）"""
    for line in code_text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        # 灰色背景
        shading = p._element.get_or_add_pPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'F0F0F0',
            qn('w:val'): 'clear'
        })
        shading.append(shd)

def parse_and_convert(md_path, docx_path):
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 页面设置
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # 一级标题 #
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:]
            doc.add_heading(text, level=0)

        # 二级标题 ##
        elif line.startswith('## '):
            text = line[3:]
            doc.add_heading(text, level=1)

        # 三级标题 ###
        elif line.startswith('### '):
            text = line[4:]
            doc.add_heading(text, level=2)

        # 四级标题 ####
        elif line.startswith('#### '):
            text = line[5:]
            doc.add_heading(text, level=3)

        # 代码块 ```
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            add_code_block(doc, '\n'.join(code_lines))

        # 表格
        elif line.startswith('|') and line.strip().endswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i].rstrip())
                i += 1
            i -= 1  # 外层循环会+1

            # 解析表格
            rows_data = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                # 跳过分隔行 (|---|---|)
                if all(re.match(r'^[-:]+$', c) for c in cells):
                    continue
                rows_data.append(cells)

            if rows_data:
                num_cols = max(len(r) for r in rows_data)
                table = doc.add_table(rows=len(rows_data), cols=num_cols, style='Table Grid')
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for ri, row_data in enumerate(rows_data):
                    for ci, cell_text in enumerate(row_data):
                        if ci < num_cols:
                            cell = table.rows[ri].cells[ci]
                            # 清除默认段落
                            cell.paragraphs[0].clear()
                            run = cell.paragraphs[0].add_run(cell_text)
                            run.font.size = Pt(9)
                            # 表头加粗 + 蓝底
                            if ri == 0:
                                run.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                set_cell_shading(cell, '2F5496')
                            elif ri % 2 == 0:
                                set_cell_shading(cell, 'D6E4F0')

                doc.add_paragraph()  # 表后空行

        # 水平线 ---
        elif line == '---':
            doc.add_paragraph('─' * 60)

        # 空行
        elif line == '':
            # 跳过连续空行
            pass

        # 普通段落 / 列表项
        else:
            # 处理行内格式
            text = line

            # 无序列表
            is_bullet = False
            if re.match(r'^(\s*)-\s+', text):
                is_bullet = True
                text = re.sub(r'^(\s*)-\s+', r'\1', text)
            elif re.match(r'^(\s*)\d+\.\s+', text):
                text = re.sub(r'^(\s*)\d+\.\s+', r'\1', text)

            # 处理行内代码 `...`
            parts = re.split(r'(`[^`]+`)', text)

            if is_bullet:
                p = doc.add_paragraph(style='List Bullet')
            else:
                p = doc.add_paragraph()

            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

            for part in parts:
                if part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
                else:
                    # 处理粗体 **...**
                    bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
                    for bp in bold_parts:
                        if bp.startswith('**') and bp.endswith('**'):
                            run = p.add_run(bp[2:-2])
                            run.bold = True
                        else:
                            run = p.add_run(bp)

        i += 1

    doc.save(docx_path)
    print(f'已生成: {docx_path}')

if __name__ == '__main__':
    md_path = r'c:\Users\Administrator\Desktop\nxp竞赛\imx93_sorting_sim\项目实现详解.md'
    docx_path = r'c:\Users\Administrator\Desktop\nxp竞赛\imx93_sorting_sim\项目实现详解.docx'
    parse_and_convert(md_path, docx_path)
