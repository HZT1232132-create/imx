#!/usr/bin/env python3
"""
Generate comprehensive tutorial document for EdgeGuard-Sort project.
Output: docs/EdgeGuard-Sort_完整教程.docx
"""
import os, sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

BASE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..")

def read_file(path):
    try:
        with open(os.path.join(BASE, path), "r", encoding="utf-8") as f:
            return f.read()
    except:
        return f"[文件未找到: {path}]"

def add_code_block(doc, code, language="cpp"):
    """Add a code block with monospace formatting."""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(30, 30, 30)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    return p

def add_section(doc, title, level=1):
    return doc.add_heading(title, level=level)

def add_text(doc, text):
    p = doc.add_paragraph(text)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r + 1].cells[c].text = str(val)
    return table

def build_tutorial():
    doc = Document()

    # ---- Title ----
    title = doc.add_heading('EdgeGuard-Sort 完整教程', level=0)
    subtitle = doc.add_paragraph(
        '基于 FRDM-i.MX93 的仓储异常标签鲁棒识别与智能分拣终端\n'
        '从零到完整系统 — 适合零基础学习者'
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ============================================================
    # 第一章：项目概述
    # ============================================================
    add_section(doc, '第一章：这个项目是干什么的？', level=1)

    add_text(doc,
        '想象一个大型仓库，传送带上的包裹需要自动识别并送到正确的区域（A区、B区、C区）。'
        '每个包裹上贴有标签，包含 QR 二维码和印刷文字。\n\n'
        '本系统模拟的就是这样一个智能分拣终端的大脑，运行在 i.MX93 嵌入式芯片上。'
        '它的工作流程非常简单：'
    )

    add_text(doc, '包裹到达 → 拍照 → 读标签 → 识别包裹ID → 查规则表 → 判断分拣是否正确 → 亮灯/蜂鸣器告警 → 记录日志')

    add_text(doc,
        '听起来很简单对吧？但真实仓库里，标签经常会损坏——二维码可能磨损、印刷文字可能模糊。'
        '所以系统设计了多级"容错"机制：先用 QR 码识别，失败了就用 OCR 文字识别，'
        'OCR 读错了还能自动纠正。这就是这个项目的核心价值。'
    )

    add_section(doc, '1.1 六张测试图片说明', level=2)
    add_table(doc, ['图片', '场景', '挑战', '系统应该怎么做'],
        [['pkg001_normal.png', '正常标签', '没有挑战', 'QR 扫码成功 → 绿灯放行'],
         ['pkg002_wrong.png', '分拣错误', '货物去了错误区域', 'QR 扫码成功但区域不对 → 红灯+蜂鸣器'],
         ['pkg003_qr_damaged.png', 'QR码损坏', '二维码物理磨损', 'QR 失败 → OCR 文字识别 → 恢复成功 → 绿灯+黄灯'],
         ['pkg004_ocr_error.png', 'OCR字符混淆', '文字印错(PKG00I)', 'QR 失败 → OCR 读错 → 自动纠错 → 黄灯复核'],
         ['pkg005_unknown.png', '未知货物', '包裹不在规则库', 'QR 读到 PKG999 → 未知货物 → 红灯人工处理'],
         ['pkg006_unreadable.png', '完全不可读', '标签彻底损坏', 'QR 失败 → OCR 也失败 → 红灯拦截']]
    )

    # ============================================================
    # 第二章：环境搭建
    # ============================================================
    add_section(doc, '第二章：环境搭建（手把手）', level=1)

    add_section(doc, '2.1 安装 WSL（Windows 用户）', level=2)
    add_text(doc,
        '如果你的电脑是 Windows，推荐使用 WSL（Windows Subsystem for Linux）。\n'
        '打开 PowerShell（管理员模式），输入：\n'
        '  wsl --install -d Ubuntu\n'
        '安装完成后重启电脑，打开 Ubuntu 终端。'
    )

    add_section(doc, '2.2 安装编译工具和依赖库', level=2)
    add_code_block(doc,
        '# 更新软件包列表\n'
        'sudo apt update\n\n'
        '# 安装编译工具\n'
        'sudo apt install -y cmake g++ make pkg-config git\n\n'
        '# 安装 OpenCV（图像处理库）\n'
        'sudo apt install -y libopencv-dev\n\n'
        '# 安装 Tesseract OCR（文字识别引擎）\n'
        'sudo apt install -y libtesseract-dev libleptonica-dev tesseract-ocr-eng\n\n'
        '# 安装 OpenSSL（用于 Hash 日志）\n'
        'sudo apt install -y libssl-dev\n\n'
        '# 安装 Python 依赖（用于模型训练）\n'
        'sudo apt install -y python3-pip python3-venv\n'
        'python3 -m venv .venv\n'
        'source .venv/bin/activate\n'
        'pip install torch torchvision --index-url https://download.pytorch.org/whl/cpu\n'
        'pip install onnx onnxruntime opencv-python-headless pillow numpy'
    )

    add_section(doc, '2.3 编译项目', level=2)
    add_code_block(doc,
        '# 进入项目目录\n'
        'cd imx93_sorting_sim\n\n'
        '# 创建编译目录\n'
        'mkdir -p build && cd build\n\n'
        '# 生成 Makefile\n'
        'cmake ..\n\n'
        '# 编译（-j4 表示用 4 个 CPU 核心并行编译）\n'
        'make -j4\n\n'
        '# 编译成功后，可执行文件在 build/sorting_sim'
    )

    # ============================================================
    # 第三章：目录结构
    # ============================================================
    add_section(doc, '第三章：项目目录结构', level=1)
    add_text(doc, '这是整个项目的文件树。每个文件的作用我们会在后面的章节逐一讲解。')

    add_code_block(doc,
        'imx93_sorting_sim/\n'
        '├── main.cpp                 ← 主程序入口，整个流水线在这里\n'
        '├── CMakeLists.txt           ← CMake 编译配置\n'
        '├── include/                 ← 头文件（声明"有什么功能"）\n'
        '│   ├── input_source.h       ← 输入源接口\n'
        '│   ├── recognizer.h         ← QR码 + OCR 识别器\n'
        '│   ├── ocr_corrector.h      ← OCR 纠错器（核心创新）\n'
        '│   ├── rule_engine.h        ← 规则引擎\n'
        '│   ├── risk_engine.h        ← 风险等级定义 + 映射\n'
        '│   ├── alarm_simulator.h    ← GPIO 告警模拟\n'
        '│   ├── logger.h             ← 事件日志\n'
        '│   ├── stats.h              ← 统计管理器\n'
        '│   ├── quality_engine.h     ← V3: 标签质量评估\n'
        '│   ├── decision_engine.h    ← V3: 可信度融合决策\n'
        '│   ├── tiny_ocr_engine.h    ← V4: Tiny-ID-OCR 推理引擎\n'
        '│   ├── inference_engine.h   ← V4: 推理引擎接口\n'
        '│   ├── hash_logger.h        ← V5: Hash 链日志\n'
        '│   ├── anomaly_detector.h   ← V5: 标签异常检测\n'
        '│   ├── label_detector.h     ← V5: 标签区域检测\n'
        '│   └── process_result.h     ← 处理结果数据结构\n'
        '├── src/                     ← 实现文件（"怎么实现功能"）\n'
        '│   ├── input_source.cpp\n'
        '│   ├── recognizer.cpp\n'
        '│   ├── ocr_corrector.cpp\n'
        '│   ├── rule_engine.cpp\n'
        '│   ├── risk_engine.cpp\n'
        '│   ├── alarm_simulator.cpp\n'
        '│   ├── logger.cpp\n'
        '│   ├── stats.cpp\n'
        '│   ├── quality_engine.cpp\n'
        '│   ├── decision_engine.cpp\n'
        '│   ├── tiny_ocr_engine.cpp\n'
        '│   ├── hash_logger.cpp\n'
        '│   ├── anomaly_detector.cpp\n'
        '│   ├── label_detector.cpp\n'
        '│   └── camera_source.cpp\n'
        '├── data/                    ← 数据文件\n'
        '│   ├── rules.csv            ← 规则表（哪些包裹去哪些区）\n'
        '│   ├── demo_sequence.csv    ← 测试序列（模拟传送带）\n'
        '│   ├── expected.csv         ← 预期结果（用于自动验证）\n'
        '│   ├── images/              ← 6 张测试图片\n'
        '│   └── synthetic/           ← 900 张合成异常标签图\n'
        '├── ml/                      ← 机器学习训练脚本\n'
        '│   ├── dataset_generator.py ← 合成训练数据\n'
        '│   ├── train_tiny_id_ocr.py ← 训练 TinyOCR 模型\n'
        '│   ├── export_onnx.py       ← 导出 ONNX 模型\n'
        '│   ├── export_tflite_int8.py← 导出 TFLite INT8\n'
        '│   ├── benchmark.py         ← 基准测试\n'
        '│   └── robustness_report.py← 鲁棒性报告\n'
        '├── tools/                   ← 工具脚本\n'
        '│   ├── verify_log.py        ← Hash 链验证\n'
        '│   ├── run_ablation.py      ← 消融实验\n'
        '│   ├── make_report.py       ← 生成最终报告\n'
        '│   └── generate_synthetic_dataset.py ← 合成数据生成\n'
        '├── models/                  ← 训练好的模型文件\n'
        '│   ├── tiny_id_ocr.onnx     ← ONNX 模型 (5MB)\n'
        '│   └── tiny_id_ocr_int8.tflite ← TFLite 量化模型 (89KB)\n'
        '├── logs/                    ← 运行日志\n'
        '├── output/                  ← 带标注的输出图片\n'
        '└── bench/                   ← 基准测试报告\n'
    )

    # ============================================================
    # 第四章：核心概念
    # ============================================================
    add_section(doc, '第四章：必须先理解的几个概念', level=1)

    add_section(doc, '4.1 什么是结构体（struct）？', level=2)
    add_text(doc,
        '结构体就像是一个"信息卡片"。比如一个学生的信息包括姓名、年龄、成绩，'
        '我们可以把这些信息打包成一个 struct：'
    )
    add_code_block(doc,
        'struct Student {\n'
        '    string name;     // 姓名\n'
        '    int age;         // 年龄\n'
        '    double score;    // 成绩\n'
        '};\n\n'
        '// 使用：\n'
        'Student s;\n'
        's.name = "小明";\n'
        's.age = 12;\n'
        's.score = 95.5;'
    )
    add_text(doc,
        '在项目中，最重要的结构体是 ProcessResult（处理结果），它记录了每个包裹处理完后的全部信息：'
        '包裹ID、识别方法、目标区域、当前区域、分拣状态、风险等级、质量评分、置信度、决策动作等等。'
    )

    add_section(doc, '4.2 什么是枚举（enum）？', level=2)
    add_text(doc,
        '枚举就是"有限的几个选项"。比如交通灯只有红、黄、绿三种颜色。\n'
        '在项目中，我们用枚举来表示识别状态（QR_SUCCESS / OCR_RECOVERED / OCR_CORRECTED / UNKNOWN_PACKAGE / LABEL_ERROR）'
        '和风险等级（LEVEL_0_NORMAL ~ LEVEL_4_CRITICAL）。'
    )

    add_section(doc, '4.3 什么是 CSV 文件？', level=2)
    add_text(doc,
        'CSV 就是用逗号分隔的表格，可以用 Excel 打开。项目中用 CSV 存储规则表和测试序列。\n'
        '比如 rules.csv（规则表）的内容是：'
    )
    add_code_block(doc,
        'package_id,name,target_zone,type\n'
        'PKG001,电子元件,A,normal\n'
        'PKG002,机械零件,B,normal\n'
        'PKG003,食品包装,A,normal\n'
        'PKG004,医疗用品,C,fragile\n'
        'PKG005,测试货物,B,test'
    )
    add_text(doc, '每一行就是一个包裹的规则：PKG001 应该去 A 区，PKG002 应该去 B 区，等等。')

    # ============================================================
    # 第五章：完整流水线
    # ============================================================
    add_section(doc, '第五章：完整处理流水线（10 个阶段）', level=1)
    add_text(doc,
        'main.cpp 是整个系统的大脑。它定义了一条"流水线"，每个包裹依次经过以下阶段：'
    )

    add_table(doc, ['阶段', '名称', '做什么', '所属版本'],
        [['Stage 0', 'Quality Assessment', '评估图片质量（模糊/强光/倾斜/遮挡/破损）', 'V3'],
         ['Stage 0.5', 'Anomaly Detection', '检测标签是否有撕毁/污渍/贴纸', 'V5'],
         ['Stage 1', 'QR Recognition', '扫描二维码获取包裹ID', 'V0'],
         ['Stage 2', 'OCR / TinyOCR', 'QR失败时用文字识别兜底', 'V1/V4'],
         ['Stage 3', 'OCR Correction', '纠正OCR识别错误（O→0, I→1等）', 'V2'],
         ['Stage 4', 'Sort Judgment', '判断包裹是否去了正确的区域', 'V0'],
         ['Stage 5', 'Risk Level', '评估风险等级 Level 0~4', 'V0'],
         ['Stage 5.5', 'Decision Engine', '5因子置信度融合 + 决策动作', 'V3/V5'],
         ['Stage 6-7', 'Timing + Alarm', '计时 + GPIO声光告警', 'V0'],
         ['Stage 8-9', 'Log + Display', '记录日志+Hash链 + 屏幕显示', 'V0/V5']]
    )

    # ============================================================
    # 第六章：核心模块详解
    # ============================================================
    add_section(doc, '第六章：核心模块逐行详解', level=1)

    # 6.1 InputSource
    add_section(doc, '6.1 InputSource — 输入源（读取测试序列）', level=2)
    add_text(doc, '作用：读取 demo_sequence.csv，把每一行变成一个 InputFrame 结构体，像播放列表一样逐个取用。')
    add_text(doc, '头文件 include/input_source.h：')
    add_code_block(doc, read_file('include/input_source.h'))
    add_text(doc, '实现文件 src/input_source.cpp — 核心是 loadSequence() 函数：')
    add_code_block(doc, read_file('src/input_source.cpp'))

    # 6.2 ProcessResult
    add_section(doc, '6.2 ProcessResult — 处理结果（最重要的数据结构）', level=2)
    add_text(doc, '作用：存储一个包裹被处理后的全部信息。这是整个系统中传递数据的"信封"。')
    add_code_block(doc, read_file('include/process_result.h'))

    # 6.3 Recognizer
    add_section(doc, '6.3 Recognizer — QR码 + OCR 识别器', level=2)
    add_text(doc,
        '作用：这是系统的"眼睛"。包含两个类：\n'
        '1. PrimaryRecognizer：用 OpenCV 内置的 QRCodeDetector 扫描二维码\n'
        '2. OCRRecognizer：用 Tesseract OCR 引擎识别图片中的文字\n\n'
        '核心技巧：\n'
        '- extractPackageId() 用正则表达式 PKG[0-9OIL]{3,} 提取包裹ID\n'
        '- 优先选纯数字后缀的候选（如 PKG003 优于 PKGO03）\n'
        '- 多 ROI 扫描：在图片的 5 个不同区域分别尝试 OCR，哪个先找到就用哪个\n'
        '- 预处理：转灰度 → 放大 → 自适应二值化'
    )
    add_code_block(doc, read_file('include/recognizer.h'))
    add_text(doc, '实现文件（关键代码）：')
    add_code_block(doc, read_file('src/recognizer.cpp'))

    # 6.4 OCRCorrector
    add_section(doc, '6.4 OCRCorrector — OCR纠错器（核心创新）', level=2)
    add_text(doc,
        '为什么需要纠错？因为 OCR 引擎经常把字符认错，比如：\n'
        '- 字母 O 被识别成数字 0\n'
        '- 字母 I 被识别成数字 1\n'
        '- 字符粘连导致丢字或多字\n\n'
        '例如：真实 ID 是 PKG001，但 OCR 读成了 PKGOOI\n\n'
        '纠错三步走：\n'
        '第1步：清洗（去除非字母数字字符，转大写）→ 精确匹配规则表\n'
        '第2步：归一化混淆字符（O→0, I/L→1）→ 再精确匹配\n'
        '第3步（前两步失败才走）：编辑距离搜索，找最接近的规则ID\n'
        '  - 短ID(≤8字符)允许1个编辑距离\n'
        '  - 长ID允许2个编辑距离\n'
        '  - 如果多个候选距离相同，拒绝纠错（防止猜错）'
    )
    add_code_block(doc, read_file('include/ocr_corrector.h'))
    add_code_block(doc, read_file('src/ocr_corrector.cpp'))

    # 6.5 RuleEngine
    add_section(doc, '6.5 RuleEngine — 规则引擎', level=2)
    add_text(doc, '作用：加载 rules.csv，建立 包裹ID → 规则 的映射表。提供三个查询：hasPackage(id)、getTargetZone(id)、getRules()。')
    add_code_block(doc, read_file('src/rule_engine.cpp'))

    # 6.6 RiskEngine
    add_section(doc, '6.6 RiskEngine — 风险引擎', level=2)
    add_text(doc,
        '作用：把识别状态 + 分拣状态 映射为 5 级风险等级。\n\n'
        '风险等级定义：\n'
        '  LEVEL_0_NORMAL   — 正常：QR扫码成功 + 正确分拣\n'
        '  LEVEL_1_LOW      — 低风险：OCR恢复成功（QR坏了但OCR修正了）\n'
        '  LEVEL_2_MEDIUM   — 中风险：OCR纠错成功（OCR读错但算法修正了）\n'
        '  LEVEL_3_HIGH     — 高风险：未知货物 / 标签完全不可读\n'
        '  LEVEL_4_CRITICAL — 危急：分拣错误（货物去了错误区域！）\n\n'
        '映射逻辑：错误分拣 > 标签异常 > OCR修正 > OCR恢复 > 正常'
    )
    add_code_block(doc, read_file('src/risk_engine.cpp'))

    # 6.7 AlarmSimulator
    add_section(doc, '6.7 AlarmSimulator — 告警模拟器', level=2)
    add_text(doc,
        '作用：模拟 i.MX93 芯片的 GPIO 引脚控制 LED 灯和蜂鸣器。\n\n'
        '告警规则：\n'
        '  LEVEL_0: 绿灯亮\n'
        '  LEVEL_1: 绿灯+黄灯亮\n'
        '  LEVEL_2: 黄灯亮\n'
        '  LEVEL_3: 红灯亮\n'
        '  LEVEL_4: 红灯亮+蜂鸣器响！'
    )
    add_code_block(doc, read_file('src/alarm_simulator.cpp'))

    # 6.8 QualityEngine (V3)
    add_section(doc, '6.8 QualityEngine — 标签质量评估（V3）', level=2)
    add_text(doc,
        '作用：用 5 个指标评估图片质量，判断标签是否可信。全部用 OpenCV 实现，不需要训练模型。\n\n'
        '5 项指标：\n'
        '1. blur（模糊）：Laplacian 方差，值越低越模糊\n'
        '2. glare（强光）：高亮+低饱和像素占比\n'
        '3. angle（倾斜）：Canny+HoughLines 检测倾斜角度\n'
        '4. occlusion（遮挡）：4×4 网格低方差检测\n'
        '5. damage（破损）：网格间边缘密度标准差\n\n'
        '综合评分 = min × 0.5 + average × 0.5\n'
        '≥0.70 → GOOD | ≥0.40 → WARNING | <0.40 → BAD'
    )
    add_code_block(doc, read_file('src/quality_engine.cpp'))

    # 6.9 DecisionEngine (V3/V5)
    add_section(doc, '6.9 DecisionEngine — 可信度融合决策（V3/V5）', level=2)
    add_text(doc,
        '作用：用 5 个因子融合计算置信度，然后做出 PASS / PASS_WITH_LOG / REVIEW / BLOCK 四种决策。\n\n'
        '5 因子融合公式（V5 细化方案）：\n'
        'confidence = 0.25×质量分 + 0.25×识别分 + 0.20×规则分 + 0.15×纠正分 + 0.15×(1-异常分)\n\n'
        '决策动作：\n'
        '  PASS          — 高置信度，自动分拣\n'
        '  PASS_WITH_LOG — 中等置信度，分拣但留审计记录\n'
        '  REVIEW        — 低置信度，标记人工复核\n'
        '  BLOCK         — 极低置信度/分拣错误，拦截'
    )
    add_code_block(doc, read_file('src/decision_engine.cpp'))

    # 6.10 EventLogger + Stats
    add_section(doc, '6.10 EventLogger + StatsManager — 日志与统计', level=2)
    add_text(doc, '作用：把每个包裹的处理结果写入 events.csv，实时统计各类指标。')
    add_code_block(doc, read_file('src/logger.cpp'))
    add_code_block(doc, read_file('src/stats.cpp'))

    # 6.11 HashLogger (V5)
    add_section(doc, '6.11 HashLogger — Hash 防篡改日志（V5）', level=2)
    add_text(doc,
        '作用：用 SHA256 给每条日志计算哈希值，形成哈希链。如果日志被篡改，链条就会断裂，可以立刻检测到。\n\n'
        '原理：\n'
        '1. 每条记录 = 事件数据 + 前一条的哈希值\n'
        '2. current_hash = SHA256(事件数据 + prev_hash)\n'
        '3. 如果有人改了第3条记录的数据，第3条的 current_hash 就对不上了\n'
        '4. 第4条的 prev_hash 也变了，所以第4条也对不上\n'
        '5. 验证工具会发现链条断裂，指出被篡改的位置\n\n'
        '这就是区块链的简化版原理！'
    )
    add_code_block(doc, read_file('src/hash_logger.cpp'))

    # 6.12 AnomalyDetector (V5)
    add_section(doc, '6.12 AnomalyDetector — 标签异常检测（V5）', level=2)
    add_text(doc,
        '作用：检测标签是否有物理损坏——撕毁、污渍、贴纸覆盖。全部用 OpenCV 实现。\n\n'
        '三种检测：\n'
        '1. tear（撕毁）：局部边缘密度的标准差，突变说明有撕裂\n'
        '2. stain（污渍）：形态学开闭运算梯度，暗/亮斑点\n'
        '3. overlay（贴纸）：Laplacian 方差在网格间的差异，贴纸纹理不同'
    )
    add_code_block(doc, read_file('src/anomaly_detector.cpp'))

    # 6.13 TinyOCREngine (V4)
    add_section(doc, '6.13 TinyOCREngine — Tiny-ID-OCR 推理引擎（V4）', level=2)
    add_text(doc,
        '作用：加载训练好的 CRNN-small 模型（ONNX Runtime），用于识别包裹上的 PKGxxx ID。\n'
        '比通用的 Tesseract OCR 更小（5MB→89KB）、更快（1-5ms）、更准（专用格式）。\n\n'
        '模型训练流程：\n'
        '1. 用 dataset_generator.py 合成 5000+1000 张训练图\n'
        '2. 用 train_tiny_id_ocr.py 训练 CRNN-small 模型\n'
        '3. 用 export_onnx.py 导出 ONNX 模型\n'
        '4. 用 export_tflite_int8.py 量化为 TFLite FP16（89KB）\n\n'
        '字符集：P, K, G, 0-9, O, I, L（共 16 个字符 + CTC blank）'
    )
    add_code_block(doc, read_file('include/tiny_ocr_engine.h'))
    add_code_block(doc, read_file('src/tiny_ocr_engine.cpp'))

    # ============================================================
    # 第七章：main.cpp 完整流程
    # ============================================================
    add_section(doc, '第七章：main.cpp — 主程序完整流程', level=1)
    add_text(doc, '以下是 main.cpp 的完整代码，带详细注释。这是整个系统的大脑，把所有模块串联起来。')
    add_code_block(doc, read_file('main.cpp'))

    # ============================================================
    # 第八章：数据文件
    # ============================================================
    add_section(doc, '第八章：数据文件说明', level=1)
    add_section(doc, '8.1 rules.csv — 规则表', level=2)
    add_code_block(doc, read_file('data/rules.csv'))
    add_section(doc, '8.2 demo_sequence.csv — 测试序列', level=2)
    add_code_block(doc, read_file('data/demo_sequence.csv'))
    add_section(doc, '8.3 expected.csv — 预期结果', level=2)
    add_text(doc, '用于自动验证：运行模拟后，用 validate.py 对比 events.csv 和 expected.csv，全部匹配才算通过。')
    add_code_block(doc, read_file('data/expected.csv'))

    # ============================================================
    # 第九章：运行和验证
    # ============================================================
    add_section(doc, '第九章：如何运行和验证', level=1)

    add_section(doc, '9.1 运行模拟', level=2)
    add_code_block(doc,
        '# 进入编译目录\n'
        'cd build\n\n'
        '# 运行（交互模式，按任意键切帧，Q 退出）\n'
        './sorting_sim\n\n'
        '# 自动播放模式（每帧 1 秒）\n'
        './sorting_sim ../data/demo_sequence.csv ../data/rules.csv ../data/images 1000\n\n'
        '# 指定模式：baseline / ocr / full / tinyocr\n'
        './sorting_sim ../data/demo_sequence.csv ../data/rules.csv ../data/images 1000 full\n\n'
        '# ONNX Runtime 需要设置动态库路径\n'
        'LD_LIBRARY_PATH=../lib ./sorting_sim ...'
    )

    add_section(doc, '9.2 验证结果', level=2)
    add_code_block(doc,
        '# 用 validate.py 对比实际输出和预期结果\n'
        'cd ..\n'
        'python3 validate.py logs/events.csv data/expected.csv\n\n'
        '# 6/6 passed 表示全部通过！'
    )

    add_section(doc, '9.3 验证 Hash 链', level=2)
    add_code_block(doc,
        '# 验证日志完整性\n'
        'python3 tools/verify_log.py logs/hash_chain.csv\n\n'
        '# 模拟篡改：\n'
        '# 1. 手动修改 logs/hash_chain.csv 中的某一行\n'
        '# 2. 再次运行验证\n'
        'python3 tools/verify_log.py logs/hash_chain_tampered.csv\n\n'
        '# 输出：*** HASH BROKEN ***  篡改被检测到！'
    )

    add_section(doc, '9.4 运行消融实验', level=2)
    add_code_block(doc,
        '# 自动运行 V0-V5 全部版本并生成对比报告\n'
        'python3 tools/run_ablation.py\n\n'
        '# 输出：bench/ablation_report.csv'
    )

    # ============================================================
    # 第十章：版本演进
    # ============================================================
    add_section(doc, '第十章：V0→V5 版本演进路线', level=1)
    add_text(doc,
        '这个项目不是一步到位的，而是分了 6 个版本逐步迭代。每个版本只增加一个核心能力。'
    )

    add_table(doc, ['版本', '一句话', '增加什么', '效果'],
        [['V0', 'QR Baseline', 'QR码扫描 + 规则匹配 + 告警', '识别率 33.3%，简单标签能处理'],
         ['V1', '+ OCR 兜底', 'QR失败后自动用Tesseract OCR', '识别率 50.0%，QR坏了也能恢复'],
         ['V2', '+ OCR 纠错', '混淆字符归一化 + 编辑距离纠错', '识别率 66.7%，OCR读错也能修正'],
         ['V3', '+ Quality Gate', '5项质量评估 + 可信度融合决策', '拥有"可信度"概念，不盲目信任'],
         ['V4', '+ Tiny-ID-OCR', '专用小模型 + ONNX/TFLite量化', '5MB→89KB，1-5ms快速推理'],
         ['V5', '+ Full System', '异常检测 + Hash链日志 + 全模式报告', '完整竞赛系统，Hash防篡改']]
    )

    # ============================================================
    # 附录
    # ============================================================
    add_section(doc, '附录A：所有运行模式一览', level=1)
    add_code_block(doc,
        '# baseline — 纯 QR 码，失败就报错（V0）\n'
        './sorting_sim ... baseline\n\n'
        '# ocr — QR + Tesseract OCR（V1）\n'
        './sorting_sim ... ocr\n\n'
        '# full — QR + OCR + 纠错器 + 质量评估 + 决策引擎（V2+V3）\n'
        './sorting_sim ... full\n\n'
        '# tinyocr — QR + Tiny-ID-OCR 专用模型（V4）\n'
        './sorting_sim ... tinyocr'
    )

    add_section(doc, '附录B：常用命令速查', level=1)
    add_table(doc, ['操作', '命令'],
        [['编译', 'cd build && cmake .. && make -j4'],
         ['运行(交互)', './sorting_sim'],
         ['运行(自动)', './sorting_sim ../data/demo_sequence.csv ../data/rules.csv ../data/images 1000 full'],
         ['运行(ONNX)', 'LD_LIBRARY_PATH=../lib ./sorting_sim ... tinyocr'],
         ['验证结果', 'python3 validate.py logs/events.csv data/expected.csv'],
         ['验证Hash', 'python3 tools/verify_log.py logs/hash_chain.csv'],
         ['消融实验', 'python3 tools/run_ablation.py'],
         ['生成报告', 'python3 tools/make_report.py'],
         ['训练模型', 'python3 ml/train_tiny_id_ocr.py dataset_id_ocr runs/tiny_id_ocr_v1 30 64'],
         ['导出ONNX', 'python3 ml/export_onnx.py runs/tiny_id_ocr_v1/best.pt'],
         ['合成数据', 'python3 tools/generate_synthetic_dataset.py']]
    )

    add_section(doc, '附录C：消融实验结果', level=1)
    add_table(doc, ['版本', '识别率', '高风险事件', '平均延迟', 'Hash验证'],
        [['V0_baseline', '33.3%', '5', '24ms', 'PASS'],
         ['V1_ocr', '50.0%', '4', '56ms', 'PASS'],
         ['V2_corrector', '66.7%', '3', '44ms', 'PASS'],
         ['V3_quality', '66.7%', '3', '48ms', 'PASS'],
         ['V4_tinyocr', '66.7%', '4', '37ms', 'PASS'],
         ['V5_full', '66.7%', '3', '43ms', 'PASS']]
    )

    # ---- Save ----
    output_path = os.path.join(BASE, "docs", "EdgeGuard-Sort_完整教程.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"教程已生成: {output_path}")
    print(f"共 {len(doc.paragraphs)} 段, {len(doc.tables)} 个表格")

if __name__ == "__main__":
    build_tutorial()
